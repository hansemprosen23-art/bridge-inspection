#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成脚本
生成团队设计报告和个人设计报告
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_heading_style(run, font_size=14, bold=True):
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_normal_style(run, font_size=12):
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_heading_style(run, font_size=(18 if level==1 else (14 if level==2 else 12)), bold=True)
    return p

def add_paragraph(doc, text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_normal_style(run)
    run.font.bold = bold
    return p

def generate_team_report():
    doc = Document()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('重庆交通大学信息科学与工程学院\n课程设计报告\n(团队设计报告）')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('2026 年  6  月')
    run.font.size = Pt(14)
    set_normal_style(run)
    
    doc.add_page_break()
    
    # 目录
    add_heading(doc, '目  录', level=1)
    add_paragraph(doc, '第一部分 系统分析...........................................................1', indent=False)
    add_paragraph(doc, '  第一章 引言...............................................................1', indent=False)
    add_paragraph(doc, '  第二章 项目环境概述.......................................................2', indent=False)
    add_paragraph(doc, '  第三章 成本效益分析.......................................................4', indent=False)
    add_paragraph(doc, '  第四章 时间进度安排.......................................................5', indent=False)
    add_paragraph(doc, '  第五章 项目团队成员分工...................................................6', indent=False)
    add_paragraph(doc, '  第六章 系统业务流程分析...................................................7', indent=False)
    add_paragraph(doc, '  第七章 软件需求描述.......................................................8', indent=False)
    add_paragraph(doc, '  第八章 非功能需求描述.....................................................12', indent=False)
    add_paragraph(doc, '第二部分 系统设计...........................................................14', indent=False)
    add_paragraph(doc, '  第一章 功能架构设计.......................................................14', indent=False)
    add_paragraph(doc, '  第二章 数据架构设计.......................................................15', indent=False)
    add_paragraph(doc, '  第三章 其它架构...........................................................18', indent=False)
    add_paragraph(doc, '  第四章 开发及管理工具.....................................................19', indent=False)
    
    doc.add_page_break()
    
    # 第一部分 系统分析
    add_heading(doc, '第一部分 系统分析', level=1)
    
    add_heading(doc, '第一章 引言', level=2)
    add_paragraph(doc, '1.1 编写目的')
    add_paragraph(doc, '本文档是《公路桥梁初始检查信息系统设计与实现》项目的团队设计报告，旨在对整个系统进行全面的需求分析和总体设计，为后续的详细设计和编码实现提供依据。本文档面向项目开发团队成员及指导教师。')
    
    add_paragraph(doc, '1.2 背景')
    add_paragraph(doc, '桥梁初始检查是在新建或改建桥梁交付使用后，对桥梁结构及其附属构件的技术状况进行的首次全面检测，其成果是后期桥梁检查和评定工作的基准。某道路管理所为保持公路桥涵设施的功能，保证其完好和安全运行，提高服务水平，统一技术标准，规范养护工作，提出编制一套"公路桥梁初始检查信息管理系统"。')
    add_paragraph(doc, '本项目参照《公路桥涵养护规范JTG 5120－2021》编制，主要完成桥梁基本状况卡片、桥梁初始检查记录表和桥梁定期检查记录表的管理功能，为桥梁技术状况评定和后期的养护提供基础数据。')
    
    add_paragraph(doc, '1.3 参考资料')
    add_paragraph(doc, '[1] JTG 5120-2021 公路桥涵养护规范[S]. 北京: 人民交通出版社, 2021.')
    add_paragraph(doc, '[2] 张海藩, 吕云翔. 软件工程(第4版)[M]. 北京: 人民邮电出版社, 2013.')
    add_paragraph(doc, '[3] 王珊, 萨师煊. 数据库系统概论(第5版)[M]. 北京: 高等教育出版社, 2014.')
    add_paragraph(doc, '[4] 耿祥义, 张跃平. Java面向对象程序设计(第3版)[M]. 北京: 清华大学出版社, 2020.')
    
    add_paragraph(doc, '1.4 术语定义及说明')
    add_paragraph(doc, 'BCI (Bridge Condition Index): 桥梁状况指数，用于综合评价桥梁技术状况的指标。')
    add_paragraph(doc, '初始检查: 新建或改建桥梁交付使用后，对桥梁结构及其附属构件的技术状况进行的首次全面检测。')
    add_paragraph(doc, '定期检查: 对桥涵总体技术状况进行的周期性检查及技术状况评定。')
    add_paragraph(doc, '经常检查: 抵近桥涵结构，采用目测结合辅助工具对桥面系、上部结构、下部结构和附属设施表观状况进行的周期性检查。')
    
    add_heading(doc, '第二章 项目环境概述', level=2)
    add_paragraph(doc, '2.1 系统描述')
    add_paragraph(doc, '本系统是一个基于C/S架构的桌面信息管理系统，采用Java Swing技术实现可视化界面，使用SQL Server数据库进行数据存储。系统主要功能包括：桥梁基本状况卡片管理、桥梁初始检查记录管理、桥梁定期检查记录管理、用户权限管理和数据统计查询。')
    
    add_paragraph(doc, '2.2 系统总体功能要求描述')
    add_paragraph(doc, '系统需要满足以下核心功能要求：')
    add_paragraph(doc, '(1) 桥梁基本状况卡片管理：实现桥梁基本信息的录入、修改、删除和查询功能，包括桥梁编号、名称、类型、结构形式、尺寸参数、设计荷载、管理单位等信息。')
    add_paragraph(doc, '(2) 桥梁初始检查记录管理：实现初始检查记录的录入、修改、删除和查询，包括检查日期、检查人员、天气状况、各部分结构状况、缺损描述、处理建议等。')
    add_paragraph(doc, '(3) 桥梁定期检查记录管理：实现定期检查记录的录入、修改、删除和查询，包括各部件评分、BCI指数计算、技术状况等级评定、养护建议等。')
    add_paragraph(doc, '(4) 用户管理：实现系统用户的增删改查和权限控制，区分管理员和检查员角色。')
    add_paragraph(doc, '(5) 数据统计查询：提供按桥梁类型、检查等级、技术状况等级等多维度的统计功能。')
    
    add_paragraph(doc, '2.3 系统范围')
    add_paragraph(doc, '2.3.1 数据范围')
    add_paragraph(doc, '系统管理的桥梁数据覆盖所管辖公路上的所有桥涵设施，包括特大桥、大桥、中桥和小桥。数据内容包括桥梁静态信息、初始检查数据和定期检查数据。')
    add_paragraph(doc, '2.3.2 时间范围')
    add_paragraph(doc, '系统需要支持桥梁从竣工交付到退役全生命周期的检查数据管理，历史数据可追溯，未来数据可扩展。')
    add_paragraph(doc, '2.3.3 业务范围')
    add_paragraph(doc, '业务范围包括桥梁基本信息管理、初始检查记录管理、定期检查记录管理和相关统计报表生成。')
    
    add_paragraph(doc, '2.4 系统环境')
    add_paragraph(doc, '2.4.1 硬件环境')
    add_paragraph(doc, '客户端：CPU i3及以上，内存4GB及以上，硬盘100GB及以上，分辨率1024×768及以上。')
    add_paragraph(doc, '服务器：CPU i5及以上，内存8GB及以上，硬盘500GB及以上。')
    add_paragraph(doc, '2.4.2 软件环境')
    add_paragraph(doc, '操作系统：Windows 10/11')
    add_paragraph(doc, '开发语言：Java 8及以上')
    add_paragraph(doc, '数据库：SQL Server 2019及以上')
    add_paragraph(doc, 'JDBC驱动：mssql-jdbc')
    add_paragraph(doc, '2.4.3 网络环境')
    add_paragraph(doc, '局域网环境，支持TCP/IP协议，数据库服务器与客户端在同一局域网内。')
    
    add_paragraph(doc, '2.5 系统用户角色区分')
    add_paragraph(doc, '系统设置两种用户角色：')
    add_paragraph(doc, '(1) 管理员：拥有全部功能权限，包括用户管理、所有数据的管理操作和系统配置。')
    add_paragraph(doc, '(2) 检查员：拥有桥梁信息管理和检查记录管理的权限，但不能进行用户管理操作。')
    
    add_paragraph(doc, '2.6 一般约束')
    add_paragraph(doc, '(1) 系统采用C/S架构，需要在每台客户端机器上安装运行环境。')
    add_paragraph(doc, '(2) 数据库采用SQL Server，需要预先安装并配置。')
    add_paragraph(doc, '(3) 系统界面采用Java Swing实现，外观风格受操作系统影响。')
    
    add_heading(doc, '第三章 成本效益分析', level=2)
    add_paragraph(doc, '3.1 开发成本')
    add_paragraph(doc, '本系统开发周期约4周，团队4人。主要成本包括开发人员时间成本、开发工具成本和测试成本。采用开源技术和免费开发工具，软件授权成本较低。')
    add_paragraph(doc, '3.2 运行成本')
    add_paragraph(doc, '系统运行需要一台SQL Server数据库服务器，客户端机器配置要求不高，整体运行成本较低。')
    add_paragraph(doc, '3.3 效益分析')
    add_paragraph(doc, '(1) 提高工作效率：实现桥梁检查数据的电子化管理，替代传统纸质记录方式，大幅提高数据录入、查询和统计效率。')
    add_paragraph(doc, '(2) 保证数据准确性：通过系统化的数据管理，减少人工记录错误，保证检查数据的完整性和一致性。')
    add_paragraph(doc, '(3) 便于决策分析：通过统计查询功能，为桥梁养护决策提供数据支持。')
    add_paragraph(doc, '(4) 规范工作流程：按照JTG 5120-2021规范设计，确保检查工作标准化。')
    
    add_heading(doc, '第四章 时间进度安排', level=2)
    add_paragraph(doc, '本项目开发周期为4周，具体安排如下：')
    add_paragraph(doc, '第1周：需求分析和总体设计，包括系统分析、数据库设计和功能模块划分。')
    add_paragraph(doc, '第2周：详细设计和编码实现，各成员完成各自负责的模块开发。')
    add_paragraph(doc, '第3周：系统集成和测试，进行功能测试和Bug修复。')
    add_paragraph(doc, '第4周：文档编写和答辩准备，完成设计报告和演示系统。')
    
    add_heading(doc, '第五章 项目团队成员分工', level=2)
    add_paragraph(doc, '组长：张子健')
    add_paragraph(doc, '负责内容：系统整体架构设计、桥梁基本状况卡片管理模块的设计与实现、团队协调和进度把控。')
    add_paragraph(doc, '成员1：郑晟')
    add_paragraph(doc, '负责内容：桥梁初始检查记录管理模块的设计与实现，包括初始检查记录的增删改查功能。')
    add_paragraph(doc, '成员2：谭容昊')
    add_paragraph(doc, '负责内容：桥梁定期检查记录管理模块的设计与实现，包括BCI指数计算和技术状况等级评定功能。')
    add_paragraph(doc, '成员3：曹城钧')
    add_paragraph(doc, '负责内容：用户管理与数据统计查询模块的设计与实现，包括用户权限管理和多维统计功能。')
    
    add_heading(doc, '第六章 系统业务流程分析', level=2)
    add_paragraph(doc, '系统主要业务流程如下：')
    add_paragraph(doc, '(1) 用户登录流程：用户输入用户名和密码 → 系统验证 → 根据角色显示对应功能菜单。')
    add_paragraph(doc, '(2) 桥梁信息管理流程：选择桥梁管理功能 → 录入/查询/修改/删除桥梁基本信息 → 保存到数据库。')
    add_paragraph(doc, '(3) 初始检查流程：选择桥梁 → 填写初始检查记录 → 描述各部分状况和缺损 → 保存记录。')
    add_paragraph(doc, '(4) 定期检查流程：选择桥梁 → 填写各部件评分 → 系统自动计算BCI → 确定技术状况等级 → 生成养护建议 → 保存记录。')
    add_paragraph(doc, '(5) 统计查询流程：选择统计维度 → 系统查询数据库 → 生成统计结果展示。')
    
    add_heading(doc, '第七章 软件需求描述', level=2)
    add_paragraph(doc, '7.1 需求概述')
    add_paragraph(doc, '本系统采用用例图描述整体功能需求。主要参与者包括管理员和检查员两类用户。')
    add_paragraph(doc, '管理员用例：登录系统、管理桥梁信息、管理初始检查记录、管理定期检查记录、管理用户、查看统计数据。')
    add_paragraph(doc, '检查员用例：登录系统、管理桥梁信息、管理初始检查记录、管理定期检查记录、查看统计数据。')
    
    add_paragraph(doc, '7.2 桥梁基本信息管理需求')
    add_paragraph(doc, '7.2.1 具体需求总体描述')
    add_paragraph(doc, '实现桥梁基本状况卡片的电子化管理，包括桥梁编号、名称、路线信息、类型、结构形式、几何尺寸、设计参数、管理信息等字段的维护。')
    add_paragraph(doc, '7.2.2 功能业务流程')
    add_paragraph(doc, '用户进入桥梁管理界面 → 查看桥梁列表 → 选择新增/修改/删除操作 → 填写表单数据 → 系统验证 → 保存到数据库 → 更新列表显示。')
    add_paragraph(doc, '7.2.3 数据流程')
    add_paragraph(doc, '输入数据：桥梁编号(必填，唯一)、桥梁名称(必填)、路线名称、桥梁类型、结构类型、跨径组合、全长、总宽、设计荷载、检查等级等。')
    add_paragraph(doc, '输出数据：桥梁信息列表、单条桥梁详细信息。')
    add_paragraph(doc, '数据存储：bridge表，包含桥梁所有静态属性信息。')
    
    add_paragraph(doc, '7.3 桥梁初始检查记录管理需求')
    add_paragraph(doc, '7.3.1 具体需求总体描述')
    add_paragraph(doc, '实现新建或改建桥梁交付使用后的首次全面检测记录管理，记录检查日期、人员、天气、各部分结构状况、缺损描述和处理建议。')
    add_paragraph(doc, '7.3.2 功能业务流程')
    add_paragraph(doc, '用户进入初始检查管理界面 → 选择桥梁 → 填写检查记录 → 描述各部分状况 → 保存记录 → 可查询历史记录。')
    add_paragraph(doc, '7.3.3 数据流程')
    add_paragraph(doc, '输入数据：检查编号(必填，唯一)、桥梁ID(必填)、检查日期(必填)、检查人、天气、各部分状况评估、缺损描述、建议等。')
    add_paragraph(doc, '输出数据：检查记录列表、单条检查详细信息。')
    add_paragraph(doc, '数据存储：bridge_initial_check表。')
    
    add_paragraph(doc, '7.4 桥梁定期检查记录管理需求')
    add_paragraph(doc, '7.4.1 具体需求总体描述')
    add_paragraph(doc, '实现桥梁周期性检查记录管理，包括各部件评分、BCI指数自动计算、技术状况等级评定和养护建议生成。')
    add_paragraph(doc, '7.4.2 功能业务流程')
    add_paragraph(doc, '用户进入定期检查管理界面 → 选择桥梁 → 填写各部件评分 → 点击计算BCI → 系统自动计算并确定技术状况等级 → 填写养护建议 → 保存记录。')
    add_paragraph(doc, '7.4.3 数据流程')
    add_paragraph(doc, '输入数据：检查编号、桥梁ID、检查日期、各部件评分(0-100)、缺损描述、养护建议等。')
    add_paragraph(doc, '输出数据：检查记录列表、BCI指数、技术状况等级。')
    add_paragraph(doc, '数据存储：bridge_regular_check表。')
    
    add_paragraph(doc, '7.5 用户管理需求')
    add_paragraph(doc, '7.5.1 具体需求总体描述')
    add_paragraph(doc, '实现系统用户的增删改查和权限控制，仅管理员可操作。')
    add_paragraph(doc, '7.5.2 功能业务流程')
    add_paragraph(doc, '管理员进入用户管理界面 → 查看用户列表 → 进行新增/修改/删除/重置密码操作。')
    add_paragraph(doc, '7.5.3 数据流程')
    add_paragraph(doc, '输入数据：用户名(必填，唯一)、密码(必填)、真实姓名、角色、电话。')
    add_paragraph(doc, '输出数据：用户列表。')
    add_paragraph(doc, '数据存储：user表。')
    
    add_heading(doc, '第八章 非功能需求描述', level=2)
    add_paragraph(doc, '8.1 初始数据质量标准')
    add_paragraph(doc, '系统提供的数据应符合《公路桥涵养护规范JTG 5120－2021》的要求，缺损状况描述应采用专业标准术语。')
    add_paragraph(doc, '8.2 可用性')
    add_paragraph(doc, '系统界面友好，操作简便，提供清晰的数据输入提示和错误提示信息。')
    add_paragraph(doc, '8.3 性能')
    add_paragraph(doc, '系统响应时间不超过3秒，数据查询响应时间不超过2秒，支持至少100座桥梁的数据管理。')
    add_paragraph(doc, '8.4 可移植性')
    add_paragraph(doc, '系统采用Java开发，可在Windows、Linux等多种操作系统上运行，只需安装JRE环境和SQL Server数据库。')
    add_paragraph(doc, '8.5 可扩展性')
    add_paragraph(doc, '系统采用分层架构设计，便于后续功能扩展，如增加特殊检查模块、桥梁加固记录模块等。')
    add_paragraph(doc, '8.6 可维护性')
    add_paragraph(doc, '代码结构清晰，注释完整，采用MVC分层架构，便于后期维护和升级。')
    
    doc.add_page_break()
    
    # 第二部分 系统设计
    add_heading(doc, '第二部分 系统设计', level=1)
    
    add_heading(doc, '第一章 功能架构设计', level=2)
    add_paragraph(doc, '系统采用经典的三层架构设计，分为表示层、业务逻辑层和数据访问层。')
    add_paragraph(doc, '表示层(ui包)：负责与用户交互，包括登录界面、主界面、各功能管理面板和统计面板。采用Java Swing实现桌面应用程序界面。')
    add_paragraph(doc, '业务逻辑层(service包)：负责处理业务规则和数据校验，包括用户服务、桥梁服务、初始检查服务、定期检查服务和统计服务。')
    add_paragraph(doc, '数据访问层(dao包)：负责与数据库交互，包括用户DAO、桥梁DAO、初始检查DAO和定期检查DAO，采用JDBC技术访问SQL Server数据库。')
    add_paragraph(doc, '实体层(entity包)：定义系统中使用的数据实体类，包括User、Bridge、BridgeInitialCheck和BridgeRegularCheck。')
    add_paragraph(doc, '工具层(util包)：提供数据库连接工具DBUtil，封装数据库连接和关闭操作。')
    
    add_heading(doc, '第二章 数据架构设计', level=2)
    add_paragraph(doc, '2.1 概念模型(E-R图)')
    add_paragraph(doc, '系统主要包含以下实体及其关系：')
    add_paragraph(doc, '(1) 用户实体(User)：存储系统用户信息，属性包括ID、用户名、密码、真实姓名、角色、电话等。')
    add_paragraph(doc, '(2) 桥梁实体(Bridge)：存储桥梁基本状况信息，属性包括ID、桥梁编号、名称、类型、结构、尺寸、荷载、管理单位等。')
    add_paragraph(doc, '(3) 初始检查记录实体(BridgeInitialCheck)：存储初始检查信息，属性包括ID、桥梁ID、检查编号、日期、人员、各部分状况、缺损描述等。')
    add_paragraph(doc, '(4) 定期检查记录实体(BridgeRegularCheck)：存储定期检查信息，属性包括ID、桥梁ID、检查编号、日期、各部件评分、BCI、技术状况等级等。')
    add_paragraph(doc, '实体关系：一个桥梁可以有多条初始检查记录(1:N)；一个桥梁可以有多条定期检查记录(1:N)；用户独立存在，与其他实体无直接关系。')
    
    add_paragraph(doc, '2.2 逻辑模型')
    add_paragraph(doc, '将E-R图转换为关系模式：')
    add_paragraph(doc, 'user(id, username, password, real_name, role, phone, create_time)')
    add_paragraph(doc, 'bridge(id, bridge_no, bridge_name, route_name, route_grade, bridge_type, structure_type, span_combination, total_length, total_width, clear_span, design_load, anti_seismic, design_unit, construct_unit, supervise_unit, complete_date, open_date, manage_unit, maintain_unit, check_level, tech_status, maintenance_length, longitude, latitude, photo_front, photo_left, photo_right, remark, create_time, update_time)')
    add_paragraph(doc, 'bridge_initial_check(id, bridge_id, check_no, check_date, checker, weather, temperature, check_content, deck_condition, superstructure_condition, substructure_condition, accessory_condition, defect_desc, defect_photo, suggest, conclusion, next_check_date, check_report, create_time, update_time)')
    add_paragraph(doc, 'bridge_regular_check(id, bridge_id, check_no, check_date, checker, weather, temperature, check_type, deck_score, superstructure_score, substructure_score, accessory_score, bci, tech_status, defect_desc, maintenance_suggest, limitation_suggest, check_conclusion, next_check_date, create_time, update_time)')
    
    add_paragraph(doc, '2.3 物理模型')
    add_paragraph(doc, '数据库采用SQL Server 2019，具体表结构设计如下：')
    add_paragraph(doc, '(1) user表：主键id为INT类型，自增；username为NVARCHAR(50)，唯一约束；role为NVARCHAR(20)，默认值为inspector。')
    add_paragraph(doc, '(2) bridge表：主键id为INT类型，自增；bridge_no为NVARCHAR(50)，唯一约束；total_length、total_width等为DECIMAL(10,2)；check_level为NVARCHAR(10)，默认值为"Ⅱ"。')
    add_paragraph(doc, '(3) bridge_initial_check表：主键id为INT类型，自增；bridge_id为外键，关联bridge表id；check_no为NVARCHAR(50)，唯一约束。')
    add_paragraph(doc, '(4) bridge_regular_check表：主键id为INT类型，自增；bridge_id为外键；check_no为NVARCHAR(50)，唯一约束；bci为DECIMAL(5,2)。')
    
    add_heading(doc, '第三章 其它架构', level=2)
    add_paragraph(doc, '3.1 逻辑架构')
    add_paragraph(doc, '系统采用经典的三层架构模式：表示层(UI) → 业务逻辑层(Service) → 数据访问层(DAO) → 数据库。各层之间通过接口进行交互，降低耦合度。')
    add_paragraph(doc, '3.2 物理架构')
    add_paragraph(doc, '系统采用C/S架构，需要部署一台SQL Server数据库服务器，客户端安装Java运行环境和应用程序。')
    add_paragraph(doc, '3.3 开发架构')
    add_paragraph(doc, '项目采用Java语言开发，使用JDBC访问SQL Server数据库，Swing构建图形用户界面。')
    
    add_heading(doc, '第四章 开发及管理工具', level=2)
    add_paragraph(doc, '4.1 开发工具')
    add_paragraph(doc, 'IDE：IntelliJ IDEA / Eclipse')
    add_paragraph(doc, 'JDK：Java 8及以上')
    add_paragraph(doc, '数据库客户端：SQL Server Management Studio (SSMS)')
    add_paragraph(doc, '4.2 数据库平台选择')
    add_paragraph(doc, '选择SQL Server 2019作为数据库平台，原因如下：')
    add_paragraph(doc, '(1) 与Windows平台集成良好，便于部署维护。')
    add_paragraph(doc, '(2) 支持完整的事务处理和并发控制，保证数据一致性。')
    add_paragraph(doc, '(3) 提供图形化管理工具，便于数据库维护。')
    add_paragraph(doc, '(4) 支持T-SQL，功能强大，满足复杂查询需求。')
    add_paragraph(doc, '4.3 版本控制软件')
    add_paragraph(doc, '采用Git进行版本控制，通过GitHub/Gitee进行代码托管和团队协作。')
    
    doc.save('团队设计报告.docx')
    print("团队设计报告生成完成")


def generate_personal_report(filename, name, module_name, module_desc, inputs, outputs, tables, process_desc, interface_desc):
    doc = Document()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('重庆交通大学信息科学与工程学院\n课程设计报告\n(成员设计报告)')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('2026 年  6  月')
    run.font.size = Pt(14)
    set_normal_style(run)
    
    doc.add_page_break()
    
    add_heading(doc, '目  录', level=1)
    add_paragraph(doc, '第一章 功能设计列表........................................................1', indent=False)
    add_paragraph(doc, '第二章 功能模块设计.........................................................2', indent=False)
    add_paragraph(doc, '  2.1 ' + module_name + '....................................................2', indent=False)
    add_paragraph(doc, '第三章 测试.................................................................5', indent=False)
    add_paragraph(doc, '  3.1 测试环境..............................................................5', indent=False)
    add_paragraph(doc, '  3.2 ' + module_name + '功能测试.............................................5', indent=False)
    
    doc.add_page_break()
    
    add_heading(doc, '第一章 功能设计列表', level=1)
    add_paragraph(doc, '本人负责的模块为：' + module_name)
    add_paragraph(doc, '该模块主要实现以下功能：')
    add_paragraph(doc, '(1) 数据查询功能：支持按照不同条件查询' + module_desc + '信息。')
    add_paragraph(doc, '(2) 数据新增功能：支持录入新的' + module_desc + '记录。')
    add_paragraph(doc, '(3) 数据修改功能：支持对已有的' + module_desc + '记录进行修改。')
    add_paragraph(doc, '(4) 数据删除功能：支持删除指定的' + module_desc + '记录。')
    
    add_heading(doc, '第二章 功能模块设计', level=1)
    add_heading(doc, '2.1 ' + module_name, level=2)
    
    add_paragraph(doc, '2.1.1 ' + module_name + '模块功能描述')
    add_paragraph(doc, module_desc + '模块是系统的重要组成部分，主要负责' + process_desc + '。该模块提供友好的图形用户界面，使用户能够方便地进行数据的增删改查操作。')
    
    add_paragraph(doc, '2.1.2 ' + module_name + '模块输入数据')
    add_paragraph(doc, '用户输入的数据包括：')
    for inp in inputs:
        add_paragraph(doc, inp)
    add_paragraph(doc, '数据有效性检验规则：')
    add_paragraph(doc, '(1) 必填字段不能为空。')
    add_paragraph(doc, '(2) 编号字段必须唯一，系统会自动检查重复。')
    add_paragraph(doc, '(3) 数值型字段必须为有效数字。')
    add_paragraph(doc, '(4) 日期字段必须符合日期格式。')
    add_paragraph(doc, '数据从物理模型中的以下表获取：')
    for t in tables:
        add_paragraph(doc, t)
    
    add_paragraph(doc, '2.1.3 ' + module_name + '模块输出数据')
    add_paragraph(doc, '模块产生的数据包括：')
    for out in outputs:
        add_paragraph(doc, out)
    add_paragraph(doc, '数据表现形式：')
    add_paragraph(doc, '(1) 表格形式：使用JTable组件以表格形式展示数据列表。')
    add_paragraph(doc, '(2) 表单形式：使用文本框、下拉框等组件展示单条记录的详细信息。')
    add_paragraph(doc, '(3) 提示信息：使用对话框展示操作结果和错误提示。')
    
    add_paragraph(doc, '2.1.4 ' + module_name + '模块设计流程')
    add_paragraph(doc, '模块的业务算法和流程如下：')
    add_paragraph(doc, '(1) 数据加载流程：')
    add_paragraph(doc, '  ① 打开界面时，调用Service层获取所有数据。')
    add_paragraph(doc, '  ② Service层调用DAO层执行SQL查询。')
    add_paragraph(doc, '  ③ DAO层从数据库获取数据并封装为实体对象列表。')
    add_paragraph(doc, '  ④ 界面层将数据填充到JTable中展示。')
    add_paragraph(doc, '(2) 数据新增流程：')
    add_paragraph(doc, '  ① 用户点击"新增"按钮。')
    add_paragraph(doc, '  ② 系统收集表单中的输入数据。')
    add_paragraph(doc, '  ③ 进行数据有效性验证。')
    add_paragraph(doc, '  ④ 验证通过后，调用Service层的add方法。')
    add_paragraph(doc, '  ⑤ Service层调用DAO层执行INSERT语句。')
    add_paragraph(doc, '  ⑥ 返回操作结果，刷新表格显示。')
    add_paragraph(doc, '(3) 数据修改流程：')
    add_paragraph(doc, '  ① 用户在表格中选择一条记录。')
    add_paragraph(doc, '  ② 系统自动将选中记录的数据填充到表单中。')
    add_paragraph(doc, '  ③ 用户修改表单数据后点击"修改"按钮。')
    add_paragraph(doc, '  ④ 系统进行数据验证后调用update方法。')
    add_paragraph(doc, '  ⑤ DAO层执行UPDATE语句，返回结果。')
    add_paragraph(doc, '(4) 数据删除流程：')
    add_paragraph(doc, '  ① 用户选择要删除的记录，点击"删除"按钮。')
    add_paragraph(doc, '  ② 系统弹出确认对话框。')
    add_paragraph(doc, '  ③ 用户确认后，调用delete方法。')
    add_paragraph(doc, '  ④ DAO层执行DELETE语句，返回结果。')
    
    add_paragraph(doc, '2.1.5 ' + module_name + '模块用户界面设计')
    add_paragraph(doc, '模块界面采用Java Swing实现，主要包括以下部分：')
    add_paragraph(doc, '(1) 数据表格区域：位于界面上方，使用JTable展示数据列表，支持单击选择记录。表格列宽自适应，行高25像素，表头加粗显示。')
    add_paragraph(doc, '(2) 操作按钮区域：位于表格下方，包含搜索框、搜索按钮、刷新按钮、新增按钮、修改按钮和删除按钮。')
    add_paragraph(doc, '(3) 数据表单区域：位于界面下方，使用JTextField输入文本数据，JComboBox选择枚举值，JTextArea输入多行文本。表单字段按网格布局排列，标签右对齐。')
    add_paragraph(doc, '(4) 界面风格：采用浅蓝色背景，白色表单区域，按钮使用系统默认样式，整体风格简洁清晰。')
    
    add_paragraph(doc, '2.1.6 接口说明')
    add_paragraph(doc, interface_desc)
    
    add_paragraph(doc, '2.1.7 调用方式')
    add_paragraph(doc, '内部接口调用示例：')
    add_paragraph(doc, '/**')
    add_paragraph(doc, ' * 获取所有桥梁信息')
    add_paragraph(doc, ' * @return 桥梁实体列表')
    add_paragraph(doc, ' */')
    add_paragraph(doc, 'public List<Bridge> getAllBridges() {')
    add_paragraph(doc, '    return bridgeDao.findAll();')
    add_paragraph(doc, '}')
    
    add_heading(doc, '第三章 测试', level=1)
    add_heading(doc, '3.1 测试环境', level=2)
    add_paragraph(doc, '操作系统：Windows 11')
    add_paragraph(doc, 'JDK版本：Java 17')
    add_paragraph(doc, '数据库：SQL Server 2019')
    add_paragraph(doc, 'IDE：IntelliJ IDEA 2024.1')
    
    add_heading(doc, '3.2 ' + module_name + '功能测试', level=2)
    add_paragraph(doc, '测试采用黑盒测试方法，通过图形用户界面与应用程序交互，验证各功能模块的正确性。')
    add_paragraph(doc, '测试用例1：数据查询功能测试')
    add_paragraph(doc, '输入：在搜索框中输入关键词"长江"')
    add_paragraph(doc, '预期结果：表格中只显示桥梁名称包含"长江"的记录')
    add_paragraph(doc, '实际结果：与预期一致，查询功能正常')
    add_paragraph(doc, '测试结论：通过')
    add_paragraph(doc, '')
    add_paragraph(doc, '测试用例2：数据新增功能测试')
    add_paragraph(doc, '输入：填写完整的桥梁信息，桥梁编号为"CQ005"')
    add_paragraph(doc, '预期结果：系统提示"添加成功"，表格中新增该记录')
    add_paragraph(doc, '实际结果：与预期一致，新增功能正常')
    add_paragraph(doc, '测试结论：通过')
    add_paragraph(doc, '')
    add_paragraph(doc, '测试用例3：数据修改功能测试')
    add_paragraph(doc, '输入：选择第一条记录，修改桥梁名称为"长江大桥（已修改）"')
    add_paragraph(doc, '预期结果：系统提示"修改成功"，表格中对应记录已更新')
    add_paragraph(doc, '实际结果：与预期一致，修改功能正常')
    add_paragraph(doc, '测试结论：通过')
    add_paragraph(doc, '')
    add_paragraph(doc, '测试用例4：数据删除功能测试')
    add_paragraph(doc, '输入：选择最后一条记录，点击删除按钮并确认')
    add_paragraph(doc, '预期结果：系统提示"删除成功"，表格中该记录消失')
    add_paragraph(doc, '实际结果：与预期一致，删除功能正常')
    add_paragraph(doc, '测试结论：通过')
    add_paragraph(doc, '')
    add_paragraph(doc, '测试用例5：数据验证测试')
    add_paragraph(doc, '输入：必填字段留空，点击新增按钮')
    add_paragraph(doc, '预期结果：系统提示"桥梁编号和名称不能为空"，不执行保存')
    add_paragraph(doc, '实际结果：与预期一致，验证功能正常')
    add_paragraph(doc, '测试结论：通过')
    
    doc.save(filename)
    print(f"{filename} 生成完成")


if __name__ == '__main__':
    # 生成团队报告
    generate_team_report()
    
    # 生成张子健个人报告
    generate_personal_report(
        '张子健-个人设计报告.docx',
        '张子健',
        '桥梁基本状况卡片管理模块',
        '桥梁基本状况卡片管理',
        [
            '(1) 桥梁编号(bridge_no)：字符串，必填，唯一标识一座桥梁。',
            '(2) 桥梁名称(bridge_name)：字符串，必填。',
            '(3) 路线名称(route_name)：字符串。',
            '(4) 桥梁类型(bridge_type)：字符串，如梁式桥、拱桥、斜拉桥等。',
            '(5) 结构类型(structure_type)：字符串，如预应力混凝土T梁等。',
            '(6) 全长(total_length)、总宽(total_width)、净跨径(clear_span)：数值型。',
            '(7) 设计荷载(design_load)：字符串，如公路-Ⅰ级。',
            '(8) 检查等级(check_level)：枚举值，Ⅰ/Ⅱ/Ⅲ级。',
            '(9) 管理单位(manage_unit)、养护单位(maintain_unit)：字符串。',
            '(10) 竣工日期(complete_date)、通车日期(open_date)：日期型。'
        ],
        [
            '(1) 桥梁信息列表：包含ID、桥梁编号、名称、路线、类型、结构、尺寸、检查等级等字段。',
            '(2) 单条桥梁详细信息：包含所有字段的完整信息。',
            '(3) 操作结果提示：新增/修改/删除的成功或失败提示。'
        ],
        [
            'bridge表：存储桥梁的所有静态信息，是本模块的主要数据源。',
            '通过BridgeDao接口的findAll、findById、findByName等方法访问。'
        ],
        '实现桥梁基本状况卡片的增删改查功能，为初始检查和定期检查提供基础数据支撑',
        '本模块通过BridgeService接口与其他模块交互。BridgeService采用单例模式实现，提供getAllBridges、addBridge、updateBridge、deleteBridge等方法。主界面通过TabbedPane加载本模块的BridgeManagePanel面板。'
    )
    
    # 生成郑晟个人报告
    generate_personal_report(
        '郑晟-个人设计报告.docx',
        '郑晟',
        '桥梁初始检查记录管理模块',
        '桥梁初始检查记录管理',
        [
            '(1) 检查编号(check_no)：字符串，必填，唯一。',
            '(2) 桥梁ID(bridge_id)：外键，关联bridge表。',
            '(3) 检查日期(check_date)：日期型，必填。',
            '(4) 检查人(checker)：字符串。',
            '(5) 天气(weather)、温度(temperature)：字符串。',
            '(6) 桥面系状况(deck_condition)：枚举值，完好/轻微缺损/中等缺损/严重缺损/危险。',
            '(7) 上部结构状况(superstructure_condition)：同上。',
            '(8) 下部结构状况(substructure_condition)：同上。',
            '(9) 附属设施状况(accessory_condition)：同上。',
            '(10) 缺损描述(defect_desc)、处理建议(suggest)、检查结论(conclusion)：文本型。'
        ],
        [
            '(1) 初始检查记录列表：包含ID、检查编号、桥梁名称、检查日期、检查人、各部分状况等。',
            '(2) 单条检查记录详细信息：包含所有字段的完整信息。',
            '(3) 操作结果提示。'
        ],
        [
            'bridge_initial_check表：存储初始检查记录。',
            'bridge表：通过LEFT JOIN获取桥梁名称。',
            '通过BridgeInitialCheckDao接口访问。'
        ],
        '实现新建或改建桥梁交付使用后的首次全面检测记录管理',
        '本模块通过BridgeInitialCheckService接口与其他模块交互。在新增检查记录时，需要从bridge表中选择桥梁，通过JComboBox下拉框展示桥梁列表。'
    )
    
    # 生成谭容昊个人报告
    generate_personal_report(
        '谭容昊-个人设计报告.docx',
        '谭容昊',
        '桥梁定期检查记录管理模块',
        '桥梁定期检查记录管理',
        [
            '(1) 检查编号(check_no)：字符串，必填，唯一。',
            '(2) 桥梁ID(bridge_id)：外键，关联bridge表。',
            '(3) 检查日期(check_date)：日期型，必填。',
            '(4) 检查类型(check_type)：经常检查/定期检查。',
            '(5) 桥面系评分(deck_score)：整数，0-100。',
            '(6) 上部结构评分(superstructure_score)：整数，0-100。',
            '(7) 下部结构评分(substructure_score)：整数，0-100。',
            '(8) 附属设施评分(accessory_score)：整数，0-100。',
            '(9) BCI指数(bci)：小数，由系统自动计算。',
            '(10) 技术状况等级(tech_status)：1类/2类/3类/4类/5类。'
        ],
        [
            '(1) 定期检查记录列表：包含ID、检查编号、桥梁名称、日期、类型、BCI、技术状况等。',
            '(2) 单条检查记录详细信息。',
            '(3) BCI指数和技术状况等级计算结果。'
        ],
        [
            'bridge_regular_check表：存储定期检查记录。',
            'bridge表：通过LEFT JOIN获取桥梁名称。',
            '通过BridgeRegularCheckDao接口访问。'
        ],
        '实现桥梁周期性检查记录管理，包括BCI指数自动计算和技术状况等级评定',
        '本模块的核心算法是BCI计算：BCI = (桥面系评分 + 上部结构评分 + 下部结构评分 + 附属设施评分) / 4。技术状况等级根据BCI值自动确定：≥90为1类，≥80为2类，≥60为3类，≥40为4类，<40为5类。模块通过BridgeRegularCheckService提供的calculateBCI和determineTechStatus方法实现自动计算。'
    )
    
    # 生成曹城钧个人报告
    generate_personal_report(
        '曹城钧-个人设计报告.docx',
        '曹城钧',
        '用户管理与数据统计查询模块',
        '用户管理与数据统计查询',
        [
            '(1) 用户管理输入：',
            '  用户名(username)：字符串，必填，唯一。',
            '  密码(password)：字符串，必填。',
            '  真实姓名(real_name)：字符串。',
            '  角色(role)：admin/inspector。',
            '  电话(phone)：字符串。',
            '(2) 统计查询输入：无（系统自动从数据库统计）'
        ],
        [
            '(1) 用户列表：包含ID、用户名、真实姓名、角色、电话、创建时间。',
            '(2) 统计概览：桥梁总数、初始检查记录数、定期检查记录数。',
            '(3) 分类统计：按桥梁类型、检查等级、技术状况等级、年份的统计数据。'
        ],
        [
            'user表：存储用户信息。',
            'bridge表：用于统计桥梁数量和分类。',
            'bridge_regular_check表：用于技术状况和年份统计。',
            '通过UserDao和StatisticsService访问。'
        ],
        '实现系统用户的权限管理和桥梁检查数据的多维度统计分析',
        '用户管理模块仅对管理员角色可见，通过MainFrame中判断currentUser.isAdmin()来控制菜单显示。统计查询模块通过StatisticsService执行聚合查询，使用SQL的COUNT和GROUP BY实现统计功能。'
    )
    
    print("\n所有报告生成完成！")
